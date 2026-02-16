"""
Form Filler Utilities
======================
Programmatic form-filling for PDF and DOCX documents.
Handles text fields, checkboxes (unicode, SDT, legacy, bracket),
radio buttons, and table cells.
"""

import re
import shutil
from typing import Dict, Optional


# ─────────────────────────────────────────────────────────────
#  PUBLIC API — called by route handlers
# ─────────────────────────────────────────────────────────────

def fill_pdf_form(src_path: str, dst_path: str, value_map: Dict[str, str]) -> None:
    """
    Fill a PDF form using PyMuPDF.
    Handles ALL widget types:
      - Text fields / combo boxes / list boxes → write string value
      - Checkboxes → check if label matches an autofill value
      - Radio buttons → select matching option
    """
    import fitz

    pdf = fitz.open(src_path)
    lower_map = {k.lower().strip(): v for k, v in value_map.items()}

    # Build a flat set of all values (lowered) for checkbox matching
    all_values_lower = {v.lower().strip() for v in value_map.values() if v}

    for page in pdf:
        # Collect text blocks for label matching
        text_blocks = _extract_text_blocks(page)

        for widget in page.widgets():
            ft = widget.field_type

            # ============= TEXT / COMBO / LIST FIELDS =============
            if ft in (fitz.PDF_WIDGET_TYPE_TEXT,
                      fitz.PDF_WIDGET_TYPE_COMBOBOX,
                      fitz.PDF_WIDGET_TYPE_LISTBOX):

                matched_value = _match_widget_to_value(
                    widget, text_blocks, lower_map,
                )
                if matched_value is not None:
                    widget.field_value = str(matched_value)
                    try:
                        widget.update()
                    except Exception:
                        pass

            # ============= CHECKBOX FIELDS =============
            elif ft == fitz.PDF_WIDGET_TYPE_CHECKBOX:
                _fill_pdf_checkbox(widget, text_blocks, lower_map, all_values_lower)

            # ============= RADIO BUTTON FIELDS =============
            elif ft == fitz.PDF_WIDGET_TYPE_RADIOBUTTON:
                _fill_pdf_radio(widget, text_blocks, lower_map, all_values_lower)

    pdf.save(dst_path)
    pdf.close()


def fill_docx_form(src_path: str, dst_path: str, value_map: Dict[str, str]) -> None:
    """
    Fill a DOCX form comprehensively:
      1. Text fields in table cells (label in one cell → value in next)
      2. Unicode checkboxes ☐ → ☑ in table cells and paragraphs
      3. SDT content controls (structured document tags)
      4. Legacy form field checkboxes (ffData)
      5. Inline placeholders (Field: __________)
      6. Bracket checkboxes [ ] → [X]
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(src_path)
    lower_map = {k.lower().strip(): v for k, v in value_map.items()}
    all_values_lower = {v.lower().strip() for v in value_map.values() if v}

    # ========== Strategy 1: Fill table cells (label → adjacent value cell) ==========
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for i, cell in enumerate(cells):
                cell_text = cell.text.strip()
                cell_lower = cell_text.lower().rstrip(":").strip()

                # Skip very long cells (likely paragraphs, not labels)
                if len(cell_lower) > 80:
                    continue

                matched_value = None
                if cell_lower in lower_map:
                    matched_value = lower_map[cell_lower]
                else:
                    # Try without parenthetical content
                    cell_no_parens = re.sub(r"\([^)]*\)", "", cell_lower).strip()
                    if cell_no_parens and cell_no_parens in lower_map:
                        matched_value = lower_map[cell_no_parens]

                if matched_value is None:
                    for field_key, val in lower_map.items():
                        if _fuzzy_field_match(cell_lower, field_key):
                            matched_value = val
                            break

                if matched_value is not None and i + 1 < len(cells):
                    target_cell = cells[i + 1]
                    target_text = target_cell.text.strip()
                    if not target_text or _is_placeholder(target_text):
                        _set_cell_text(target_cell, str(matched_value))

        # Also try: label and value in SAME cell separated by colon or underscores
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if ":" in cell_text:
                    parts = cell_text.split(":", 1)
                    label_part = parts[0].strip().lower()
                    value_part = parts[1].strip() if len(parts) > 1 else ""
                    if _is_placeholder(value_part) or not value_part:
                        matched_value = None
                        if label_part in lower_map:
                            matched_value = lower_map[label_part]
                        else:
                            for field_key, val in lower_map.items():
                                if _fuzzy_field_match(label_part, field_key):
                                    matched_value = val
                                    break
                        if matched_value is not None:
                            # Rewrite cell: "Label: <value>"
                            new_text = parts[0].strip() + ": " + str(matched_value)
                            for para in cell.paragraphs:
                                if ":" in para.text:
                                    if para.runs:
                                        full = para.text
                                        idx = full.index(":")
                                        before = full[:idx + 1]
                                        para.runs[0].text = before + " " + str(matched_value)
                                        for run in para.runs[1:]:
                                            run.text = ""
                                    break

    # ========== Strategy 2: Unicode checkboxes ☐ → ☑ in ALL elements ==========
    _fill_unicode_checkboxes_in_tables(doc, lower_map, all_values_lower)
    _fill_unicode_checkboxes_in_paragraphs(doc, lower_map, all_values_lower)

    # ========== Strategy 3: SDT content control checkboxes ==========
    _fill_sdt_checkboxes(doc, lower_map, all_values_lower)

    # ========== Strategy 4: Legacy form field checkboxes ==========
    _fill_legacy_checkboxes(doc, lower_map, all_values_lower)

    # ========== Strategy 5: Inline placeholders in paragraphs ==========
    for para in doc.paragraphs:
        para_text = para.text
        for field_key, val in lower_map.items():
            pattern = re.compile(
                re.escape(field_key) + r"\s*[:]\s*[_\[\]\s]{2,}",
                re.IGNORECASE,
            )
            match = pattern.search(para_text)
            if match:
                new_text = para_text[:match.start()] + field_key + ": " + str(val)
                remainder = para_text[match.end():]
                new_text += remainder
                if para.runs:
                    para.runs[0].text = new_text
                    for run in para.runs[1:]:
                        run.text = ""
                break

    # ========== Strategy 6: Bracket checkboxes [ ] → [X] ==========
    _fill_bracket_checkboxes(doc, lower_map, all_values_lower)

    doc.save(dst_path)


# ─────────────────────────────────────────────────────────────
#  PDF HELPERS
# ─────────────────────────────────────────────────────────────

def _extract_text_blocks(page) -> list:
    """Extract text blocks with bounding boxes from a PDF page."""
    import fitz
    text_blocks = []
    try:
        text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        for block in text_dict.get("blocks", []):
            if block.get("type") == 0:
                for line in block.get("lines", []):
                    line_text = "".join(
                        span.get("text", "") for span in line.get("spans", [])
                    ).strip()
                    if line_text:
                        text_blocks.append({
                            "text": line_text,
                            "bbox": fitz.Rect(line["bbox"]),
                        })
    except Exception:
        pass
    return text_blocks


def _match_widget_to_value(widget, text_blocks, lower_map) -> Optional[str]:
    """Try to match a PDF widget to an autofill value using multiple strategies."""
    # Strategy 1: Direct match on widget.field_name
    wname = (widget.field_name or "").strip()
    wname_lower = wname.lower()
    if wname_lower in lower_map:
        return lower_map[wname_lower]

    # Strategy 1b: Strip common suffixes/prefixes from widget name
    wname_clean = re.sub(r"[\[\(]\d+[\]\)]$", "", wname_lower).strip()
    wname_clean = re.sub(r"_\d+$", "", wname_clean).strip()
    wname_clean = wname_clean.rstrip(":").strip()
    if wname_clean and wname_clean in lower_map:
        return lower_map[wname_clean]

    # Strategy 1c: Normalize underscores to spaces and retry
    wname_spaced = re.sub(r"[_\-]", " ", wname_clean).strip()
    if wname_spaced and wname_spaced in lower_map:
        return lower_map[wname_spaced]

    # Strategy 2: Fuzzy match on widget.field_name
    for field_key, val in lower_map.items():
        if _fuzzy_field_match(wname_lower, field_key):
            return val
    if wname_clean and wname_clean != wname_lower:
        for field_key, val in lower_map.items():
            if _fuzzy_field_match(wname_clean, field_key):
                return val

    # Strategy 3: Match using nearest text label (primary — left/above)
    if text_blocks:
        label = _find_nearest_label(widget.rect, text_blocks)
        if label:
            label_lower = label.lower().strip().rstrip(":").strip()
            if label_lower in lower_map:
                return lower_map[label_lower]
            # Try without parenthetical content
            label_no_parens = re.sub(r"\([^)]*\)", "", label_lower).strip()
            if label_no_parens and label_no_parens in lower_map:
                return lower_map[label_no_parens]
            for field_key, val in lower_map.items():
                if _fuzzy_field_match(label_lower, field_key):
                    return val

        # Strategy 4: Try multiple nearby labels
        nearby = _find_nearby_texts(widget.rect, text_blocks, max_dist=250.0, limit=5)
        for candidate in nearby:
            cand_lower = candidate.lower().strip().rstrip(":").strip()
            if cand_lower in lower_map:
                return lower_map[cand_lower]
            # Also try without parenthetical
            cand_no_parens = re.sub(r"\([^)]*\)", "", cand_lower).strip()
            if cand_no_parens and cand_no_parens in lower_map:
                return lower_map[cand_no_parens]
        for candidate in nearby:
            cand_lower = candidate.lower().strip().rstrip(":").strip()
            for field_key, val in lower_map.items():
                if _fuzzy_field_match(cand_lower, field_key):
                    return val

    return None


def _fill_pdf_checkbox(widget, text_blocks, lower_map, all_values_lower) -> None:
    """Fill a PDF checkbox widget."""
    label = _find_nearest_label(widget.rect, text_blocks, max_dist=150.0)
    if not label:
        return

    label_clean = label.lower().strip()
    for sym in ["\u2610", "\u2611", "\u2612", "\u2713", "\u2717"]:
        label_clean = label_clean.replace(sym, "")
    label_clean = label_clean.strip()

    if not label_clean:
        return

    should_check = False
    _YESNO_WORDS = {"yes", "no", "true", "false"}

    # --- Check 1: SPECIFIC label matches an autofill value ---
    if label_clean not in _YESNO_WORDS:
        safe_values = {v for v in all_values_lower if v not in _YESNO_WORDS}
        if label_clean in safe_values:
            should_check = True
        if not should_check:
            for val in safe_values:
                if label_clean in val or val in label_clean:
                    should_check = True
                    break
        if not should_check:
            for val in safe_values:
                if _fuzzy_field_match(label_clean, val):
                    should_check = True
                    break

    # --- Check 2: Yes/No checkbox — match against QUESTION context ---
    if not should_check and label_clean in _YESNO_WORDS:
        nearby_texts = _find_nearby_texts(widget.rect, text_blocks, max_dist=400.0, limit=15)
        for question_text in nearby_texts:
            q_lower = question_text.lower().strip()
            for sym in ["\u2610", "\u2611", "\u2612"]:
                q_lower = q_lower.replace(sym, "")
            q_lower = q_lower.strip()
            if not q_lower or q_lower in _YESNO_WORDS:
                continue

            for field_key, val in lower_map.items():
                if _fuzzy_field_match(q_lower, field_key):
                    val_lower = val.lower().strip()
                    yes_answers = {"yes", "true", "y", "1", "checked", "approved"}
                    no_answers = {"no", "false", "n", "0", "unchecked", "not approved"}
                    if label_clean in ("yes", "true") and val_lower in yes_answers:
                        should_check = True
                    elif label_clean in ("no", "false") and val_lower in no_answers:
                        should_check = True
                    break
            if should_check:
                break

    if should_check:
        widget.field_value = True
        try:
            widget.update()
        except Exception:
            pass


def _fill_pdf_radio(widget, text_blocks, lower_map, all_values_lower) -> None:
    """Fill a PDF radio button."""
    label = _find_nearest_label(widget.rect, text_blocks, max_dist=150.0)
    if not label:
        return

    label_clean = label.lower().strip()
    for sym in ["\u2610", "\u2611", "\u2612"]:
        label_clean = label_clean.replace(sym, "")
    label_clean = label_clean.strip()

    _YESNO = {"yes", "no", "true", "false"}
    safe_values = {v for v in all_values_lower if v not in _YESNO}

    if label_clean in safe_values:
        widget.field_value = True
        try:
            widget.update()
        except Exception:
            pass
    else:
        for val in safe_values:
            if _fuzzy_field_match(label_clean, val):
                widget.field_value = True
                try:
                    widget.update()
                except Exception:
                    pass
                break


def _find_nearby_texts(widget_rect, text_blocks, max_dist: float = 350.0,
                       limit: int = 5) -> list:
    """Find multiple nearby text labels for context."""
    wx_center = (widget_rect.x0 + widget_rect.x1) / 2
    wy_center = (widget_rect.y0 + widget_rect.y1) / 2

    scored = []
    for block in text_blocks:
        bbox = block["bbox"]
        bx_center = (bbox.x0 + bbox.x1) / 2
        by_center = (bbox.y0 + bbox.y1) / 2
        dx = wx_center - bx_center
        dy = wy_center - by_center
        if abs(dy) < 15:
            dist = abs(dx) * 0.5
        elif dy > 0:
            dist = (dx ** 2 + dy ** 2) ** 0.5
        else:
            dist = (dx ** 2 + dy ** 2) ** 0.5 * 2
        if dist < max_dist and block["text"].strip():
            scored.append((dist, block["text"].strip()))

    scored.sort(key=lambda x: x[0])
    return [text for _, text in scored[:limit]]


def _find_nearest_label(widget_rect, text_blocks, max_dist: float = 200.0) -> Optional[str]:
    """Find the closest text label to a PDF widget."""
    best_label = None
    best_dist = max_dist

    wx_center = (widget_rect.x0 + widget_rect.x1) / 2
    wy_center = (widget_rect.y0 + widget_rect.y1) / 2

    for block in text_blocks:
        bbox = block["bbox"]
        bx_center = (bbox.x0 + bbox.x1) / 2
        by_center = (bbox.y0 + bbox.y1) / 2
        text = block["text"].strip()

        if not text:
            continue

        dx = wx_center - bx_center
        dy = wy_center - by_center

        if abs(dy) < 18:
            if dx > 0:
                dist = abs(dx) * 0.4
            else:
                dist = abs(dx) * 0.8
        elif dy > 0:
            dist = (dx ** 2 + dy ** 2) ** 0.5 * 0.9
        else:
            dist = (dx ** 2 + dy ** 2) ** 0.5 * 2.5

        if dist < best_dist:
            best_dist = dist
            best_label = text

    return best_label


# ─────────────────────────────────────────────────────────────
#  DOCX HELPERS
# ─────────────────────────────────────────────────────────────

_UNCHECKED = "\u2610"  # ☐
_CHECKED   = "\u2611"  # ☑


def _should_check_option(option_label: str, lower_map: Dict[str, str],
                         all_values_lower: set, context: str = "") -> bool:
    """Determine whether a checkbox option should be checked."""
    opt = option_label.lower().strip()
    if not opt:
        return False

    _YESNO = {"yes", "no", "true", "false", "y", "n"}

    if opt not in _YESNO:
        safe_values = {v for v in all_values_lower if v not in _YESNO}
        if opt in safe_values:
            return True
        for val in safe_values:
            if opt in val or val in opt:
                return True
        for val in safe_values:
            if _fuzzy_field_match(opt, val):
                return True
        return False

    if context:
        ctx_lower = context.lower().strip()
        for field_key, val in lower_map.items():
            if _fuzzy_field_match(ctx_lower, field_key):
                val_lower = val.lower().strip()
                yes_answers = {"yes", "true", "y", "1", "checked", "approved"}
                no_answers = {"no", "false", "n", "0", "unchecked", "not approved"}
                if opt in ("yes", "true", "y") and val_lower in yes_answers:
                    return True
                if opt in ("no", "false", "n") and val_lower in no_answers:
                    return True
                break

    return False


def _fill_unicode_checkboxes_in_tables(doc, lower_map, all_values_lower) -> None:
    """Find ☐ in table cells and replace with ☑ when matching."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_paragraph_unicode_checkboxes(
                        para, lower_map, all_values_lower,
                        context=cell.text,
                    )


def _fill_unicode_checkboxes_in_paragraphs(doc, lower_map, all_values_lower) -> None:
    """Find ☐ in top-level paragraphs and replace with ☑ when matching."""
    for para in doc.paragraphs:
        _process_paragraph_unicode_checkboxes(
            para, lower_map, all_values_lower,
            context=para.text,
        )


def _process_paragraph_unicode_checkboxes(para, lower_map, all_values_lower,
                                           context: str = "") -> None:
    """Process a single paragraph: find ☐ chars and replace with ☑."""
    full_text = para.text
    if _UNCHECKED not in full_text:
        return

    pairs = _parse_checkbox_pairs(full_text)
    if not pairs:
        return

    options_to_check = set()
    for option_label in pairs:
        if _should_check_option(option_label, lower_map, all_values_lower, context):
            options_to_check.add(option_label)

    if not options_to_check:
        return

    run_offset = 0
    for run in para.runs:
        run_text = run.text
        if _UNCHECKED in run_text:
            new_run_text = list(run_text)
            for ci, ch in enumerate(run_text):
                if ch == _UNCHECKED:
                    abs_pos = run_offset + ci
                    option = _get_option_at_position(full_text, abs_pos)
                    if option and option in options_to_check:
                        new_run_text[ci] = _CHECKED
            run.text = "".join(new_run_text)
        run_offset += len(run_text)


def _parse_checkbox_pairs(text: str) -> list:
    """Parse checkbox-label pairs from text."""
    parts = text.split(_UNCHECKED)
    labels = []
    for part in parts[1:]:
        label = part.strip()
        for sym in [_CHECKED, "\u2612"]:
            label = label.replace(sym, "")
        label = label.strip().rstrip(",;.")
        if label:
            labels.append(label.lower().strip())
    return labels


def _get_option_at_position(full_text: str, checkbox_pos: int) -> Optional[str]:
    """Get the option label for the checkbox at a given position."""
    after = full_text[checkbox_pos + 1:]
    after = after.strip()
    if not after:
        return None

    end = len(after)
    for sym in [_UNCHECKED, _CHECKED, "\u2612"]:
        idx = after.find(sym)
        if idx != -1 and idx < end:
            end = idx

    label = after[:end].strip().rstrip(",;.")
    return label.lower().strip() if label else None


def _fill_sdt_checkboxes(doc, lower_map, all_values_lower) -> None:
    """Fill structured document tag (SDT) checkboxes in the DOCX XML."""
    from lxml import etree

    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns_w14 = "http://schemas.microsoft.com/office/word/2010/wordml"
    nsmap = {"w": ns_w, "w14": ns_w14}

    body = doc.element
    sdts = body.findall(".//w:sdt", nsmap)

    for sdt in sdts:
        cb14 = sdt.find(".//w14:checkbox", nsmap)
        if cb14 is None:
            continue

        alias_el = sdt.find(".//w:sdtPr/w:alias", nsmap)
        tag_el = sdt.find(".//w:sdtPr/w:tag", nsmap)
        qn_val = f"{{{ns_w}}}val"

        alias = alias_el.get(qn_val) if alias_el is not None else ""
        tag = tag_el.get(qn_val) if tag_el is not None else ""
        label = (alias or tag or "").lower().strip()

        if not label:
            content = sdt.find(".//w:sdtContent", nsmap)
            if content is not None:
                texts = content.itertext()
                label = "".join(texts).strip().lower()

        if _should_check_option(label, lower_map, all_values_lower, ""):
            checked_el = cb14.find("w14:checked", nsmap)
            if checked_el is not None:
                checked_el.set(f"{{{ns_w14}}}val", "1")
            else:
                etree.SubElement(cb14, f"{{{ns_w14}}}checked",
                                 {f"{{{ns_w14}}}val": "1"})

            content = sdt.find(".//w:sdtContent", nsmap)
            if content is not None:
                for t_el in content.iter(f"{{{ns_w}}}t"):
                    if t_el.text and _UNCHECKED in t_el.text:
                        t_el.text = t_el.text.replace(_UNCHECKED, _CHECKED)


def _fill_legacy_checkboxes(doc, lower_map, all_values_lower) -> None:
    """Fill legacy form field checkboxes (w:ffData / w:checkBox)."""
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    nsmap = {"w": ns_w}
    qn_val = f"{{{ns_w}}}val"

    body = doc.element
    for ff in body.findall(".//w:ffData", nsmap):
        cb_el = ff.find("w:checkBox", nsmap)
        if cb_el is None:
            continue

        name_el = ff.find("w:name", nsmap)
        label = (name_el.get(qn_val) if name_el is not None else "").lower().strip()

        if _should_check_option(label, lower_map, all_values_lower, ""):
            default_el = cb_el.find("w:default", nsmap)
            if default_el is not None:
                default_el.set(qn_val, "1")
            checked_el = cb_el.find("w:checked", nsmap)
            if checked_el is not None:
                checked_el.set(qn_val, "1")
            else:
                from lxml import etree
                etree.SubElement(cb_el, f"{{{ns_w}}}checked", {qn_val: "1"})


def _fill_bracket_checkboxes(doc, lower_map, all_values_lower) -> None:
    """Fill bracket-style checkboxes: [ ] → [X]."""
    bracket_re = re.compile(r"\[\s*\]")

    def _process_runs(runs, context=""):
        full = "".join(r.text for r in runs)
        if not bracket_re.search(full):
            return
        for run in runs:
            if bracket_re.search(run.text):
                parts = bracket_re.split(run.text)
                if len(parts) > 1:
                    after_text = parts[1].strip().split("\n")[0].strip()
                    opt = after_text.lower().rstrip(",;.").strip()
                    if _should_check_option(opt, lower_map, all_values_lower, context):
                        run.text = bracket_re.sub("[X]", run.text, count=1)

    for para in doc.paragraphs:
        _process_runs(para.runs, para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_runs(para.runs, cell.text)


# ─────────────────────────────────────────────────────────────
#  SHARED MATCHING HELPERS
# ─────────────────────────────────────────────────────────────

_STOP_WORDS = {
    "a", "an", "the", "is", "are", "was", "were", "be", "been", "being",
    "have", "has", "had", "do", "does", "did", "will", "would", "could",
    "should", "may", "might", "shall", "can", "need", "dare", "ought",
    "and", "or", "but", "nor", "not", "so", "yet", "both", "either",
    "neither", "if", "then", "else", "when", "where", "while", "as",
    "at", "by", "for", "from", "in", "into", "of", "on", "to", "with",
    "it", "its", "you", "your", "we", "our", "they", "their", "he",
    "she", "his", "her", "this", "that", "these", "those", "all", "any",
    "no", "each", "every", "such", "what", "which", "who", "whom",
    "how", "than", "very", "just", "about", "also", "only",
}


def _fuzzy_field_match(text: str, field_key: str) -> bool:
    """Check if two field identifiers are semantically close enough."""
    if not text or not field_key:
        return False

    text = text.rstrip(":").strip()
    field_key = field_key.rstrip(":").strip()

    if not text or not field_key:
        return False

    # Exact substring match
    if field_key in text or text in field_key:
        return True

    # Strip parenthetical content and retry
    text_no_parens = re.sub(r"\([^)]*\)", "", text).strip()
    fk_no_parens = re.sub(r"\([^)]*\)", "", field_key).strip()
    if text_no_parens and fk_no_parens:
        if fk_no_parens in text_no_parens or text_no_parens in fk_no_parens:
            return True

    # Normalize underscores/hyphens to spaces for matching
    text_norm = re.sub(r"[_\-]", " ", text).strip()
    fk_norm = re.sub(r"[_\-]", " ", field_key).strip()
    if text_norm and fk_norm:
        if fk_norm in text_norm or text_norm in fk_norm:
            return True

    # Word-overlap matching (semantic closeness)
    words_a = set(re.findall(r"[a-z]+", text_norm)) - _STOP_WORDS
    words_b = set(re.findall(r"[a-z]+", fk_norm)) - _STOP_WORDS
    if not words_a or not words_b:
        return False
    overlap = words_a & words_b
    shorter = min(len(words_a), len(words_b))
    if shorter > 0 and len(overlap) / shorter >= 0.6:
        return True

    return False


def _is_placeholder(text: str) -> bool:
    """Check if text looks like an empty placeholder."""
    cleaned = text.strip()
    if not cleaned:
        return True
    if all(c in "_-\u2013\u2014 .[](){}|/\t\n" for c in cleaned):
        return True
    _PLACEHOLDER_STRINGS = {
        "\u2014", "-", "\u2013", "n/a", "/", "...", "___", "____",
        "________", "tbd", "enter", "fill in", "fill",
        "(enter)", "[enter]", "click here",
    }
    if cleaned.lower() in _PLACEHOLDER_STRINGS:
        return True
    if re.fullmatch(r"[_\-\u2013\u2014\s.]+", cleaned):
        return True
    return False


def _set_cell_text(cell, text: str) -> None:
    """Set text of a DOCX table cell, preserving first paragraph formatting."""
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.text = text
        for extra_p in cell.paragraphs[1:]:
            extra_p.text = ""
    else:
        cell.text = text


# ─────────────────────────────────────────────────────────────
#  FILE HELPERS
# ─────────────────────────────────────────────────────────────

def safe_remove(path: str) -> None:
    """Safely remove a single file."""
    try:
        import os
        if os.path.exists(path):
            os.remove(path)
    except Exception:
        pass


def cleanup_temp(directory: str) -> None:
    """Remove a temporary directory and all its contents."""
    try:
        import os
        if os.path.exists(directory):
            shutil.rmtree(directory, ignore_errors=True)
    except Exception:
        pass


